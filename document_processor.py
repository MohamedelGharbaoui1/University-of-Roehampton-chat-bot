# document_processor.py - Document reading and processing

import logging
from pathlib import Path
from typing import Tuple, Optional, Dict, Any
from PyPDF2 import PdfReader
from docx import Document
from config import Config

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document reading and processing operations"""
    
    @staticmethod
    def read_document(file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Read PDF or DOCX document"""
        try:
            if file_path.suffix.lower() == '.pdf':
                return DocumentProcessor._read_pdf(file_path)
            elif file_path.suffix.lower() == '.docx':
                return DocumentProcessor._read_docx(file_path)
            else:
                return None, {'error': f'Unsupported file type: {file_path.suffix}'}
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return None, {'error': str(e)}
    
    @staticmethod
    def _read_pdf(file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Read PDF file and extract metadata"""
        try:
            reader = PdfReader(str(file_path))
            text = ""
            total_pages = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            metadata = {
                'total_pages': total_pages,
                'file_size': file_path.stat().st_size,
                'file_type': 'PDF',
                'word_count': len(text.split()) if text else 0,
                'character_count': len(text),
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path.name}: {e}")
            return None, {'error': str(e)}
    
    @staticmethod
    def _read_docx(file_path: Path) -> Tuple[Optional[str], Dict[str, Any]]:
        """Read DOCX file"""
        try:
            doc = Document(str(file_path))
            text_parts = []
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text_parts.append(para_text)
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = len(doc.tables)
            for table_idx, table in enumerate(doc.tables):
                text_parts.append(f"\n--- Table {table_idx + 1} ---")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n".join(text_parts)
            
            metadata = {
                'paragraphs': paragraph_count,
                'tables': table_count,
                'file_size': file_path.stat().st_size,
                'file_type': 'Word Document',
                'word_count': len(full_text.split()) if full_text else 0,
                'character_count': len(full_text),
            }
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path.name}: {e}")
            return None, {'error': str(e)}
    
    @staticmethod
    def load_document_for_module(module_data: Dict) -> Tuple[Optional[str], Dict[str, Any], str]:
        """Enhanced document loading with support for multiple PDFs"""
        try:
            combined_content = ""
            combined_metadata = {}
            loaded_files = []
            
            if module_data['pdf_file'] == 'multiple':
                # Load all PDFs for this module
                total_pages = 0
                total_words = 0
                total_size = 0
                
                for pdf_data in module_data['all_pdfs']:
                    pdf_filename = pdf_data['pdf_file']
                    pdf_path = Path(Config.DATA_FOLDER) / pdf_filename
                    
                    if pdf_path.exists():
                        content, metadata = DocumentProcessor.read_document(pdf_path)
                        if content:
                            # Add section header for each document
                            section_header = f"\n\n{'='*60}\n📄 {pdf_data['display_name']} ({pdf_data['coursework_type']})\nFile: {pdf_data['pdf_file']}\n{'='*60}\n"
                            combined_content += section_header + content
                            
                            # Aggregate metadata
                            total_pages += metadata.get('total_pages', 0)
                            total_words += metadata.get('word_count', 0)
                            total_size += metadata.get('file_size', 0)
                            loaded_files.append(pdf_filename)
                        else:
                            logger.warning(f"Failed to load content from {pdf_filename}")
                    else:
                        logger.warning(f"File not found: {pdf_filename}")
                
                combined_metadata = {
                    'total_pages': total_pages,
                    'total_words': total_words,
                    'total_size': total_size,
                    'loaded_files': loaded_files,
                    'file_count': len(loaded_files),
                    'file_type': 'Multiple Documents'
                }
                
                if combined_content:
                    return combined_content, combined_metadata, f"Loaded {len(loaded_files)} documents for {module_data['module']}"
                else:
                    return None, combined_metadata, f"No content could be loaded for {module_data['module']}"
            
            else:
                # Load single PDF
                pdf_filename = module_data['pdf_file']
                pdf_path = Path(Config.DATA_FOLDER) / pdf_filename
                
                if not pdf_path.exists():
                    return None, {}, f"Document not found: {pdf_filename}"
                
                content, metadata = DocumentProcessor.read_document(pdf_path)
                
                if content:
                    return content, metadata, f"Loaded {pdf_filename} successfully"
                else:
                    return None, metadata, f"Failed to extract content from {pdf_filename}"
                    
        except Exception as e:
            logger.error(f"Error loading document for module: {e}")
            return None, {}, f"Error: {str(e)}"
    
    @staticmethod
    def get_document_preview(content: str, max_length: int = None) -> str:
        """Get a preview of document content"""
        if not content:
            return "No content available"
        
        max_length = max_length or Config.PREVIEW_LENGTH
        
        if len(content) <= max_length:
            return content
        
        # Find a good breaking point (end of sentence or paragraph)
        preview = content[:max_length]
        
        # Try to break at the end of a sentence
        last_period = preview.rfind('.')
        last_newline = preview.rfind('\n')
        
        break_point = max(last_period, last_newline)
        
        if break_point > max_length * 0.7:  # If we found a good break point
            preview = content[:break_point + 1]
        
        return preview + "..." if len(content) > len(preview) else preview
